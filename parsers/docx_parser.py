from __future__ import annotations

from pathlib import Path

from docx import Document

from models import Partner, ParseResult
from parsers.base import BaseParser
from parsers.table_tools import build_column_map, find_header_index, rows_to_items
from utils.text import clean_text


class DocxParser(BaseParser):
    file_format = "docx"

    def parse(self, path: Path, partner: Partner | None = None) -> ParseResult:
        partner = self.make_partner(path, partner)
        document = self.make_document(path, partner)
        result = ParseResult(document=document)
        doc = Document(path)
        result.document.raw_content = self._extract_raw_text(doc)
        for table in doc.tables:
            result.items.extend(self._parse_table(table, document))
        if not result.items:
            result.errors.append("DOCX document has no recognized price rows")
        return result.finalize()

    def _parse_table(self, table, document) -> list:
        rows = [[clean_text(cell.text) for cell in row.cells] for row in table.rows]
        header_index = find_header_index(rows)
        if header_index is None:
            return []
        columns = build_column_map(rows, header_index)
        return rows_to_items(rows, document, columns, header_index + 1)

    @staticmethod
    def _extract_raw_text(doc: Document) -> str:
        paragraphs = [clean_text(p.text) for p in doc.paragraphs if clean_text(p.text)]
        table_rows = []
        for table in doc.tables:
            for row in table.rows:
                table_rows.append(" | ".join(clean_text(cell.text) for cell in row.cells))
        return "\n".join(paragraphs + table_rows[:500])
